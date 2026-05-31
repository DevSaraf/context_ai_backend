"""
skill_docx.py
=============
Render a compiled SKILL.md (the `skill_md` string stored on a CompiledSkill)
into a polished Word .docx, returned as raw bytes for streaming download.

Pure-Python (python-docx) — no system binaries, no Node — so it runs inside a
locked-down server environment. Parses the KRAB skill_md format:

    ---
    name: <slug>
    description: "<...>"
    x-krab:
      procedure_id: 29
      version: 1
      status: verified
      autonomy_level: suggest_only
      ...
    ---

    # Title
    ## Section
    **bold lead-in** normal text
    - bullet
    1. numbered step

Only a pragmatic subset of markdown is handled — exactly what the skill
compiler emits: ATX headings (#, ##, ###), unordered list items (-, *),
ordered list items (1. 2. ...), **bold** inline spans, `inline code`, blank-line
paragraph breaks, and a YAML frontmatter block rendered as a metadata table.
"""

from __future__ import annotations

import io
import re
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT


# --------------------------------------------------------------------------- #
# Frontmatter split
# --------------------------------------------------------------------------- #
def _split_frontmatter(md: str) -> Tuple[List[str], str]:
    """Return (frontmatter_lines, body_markdown).

    Frontmatter is the block between the first pair of '---' fences. If absent,
    returns ([], full_text).
    """
    lines = md.splitlines()
    if not lines or lines[0].strip() != "---":
        return [], md
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            return lines[1:i], "\n".join(lines[i + 1 :]).lstrip("\n")
    # no closing fence — treat whole thing as body
    return [], md


def _flatten_frontmatter(fm_lines: List[str]) -> List[Tuple[str, str]]:
    """Turn YAML-ish frontmatter into flat (key, value) pairs for a table.

    Handles one level of nesting (e.g. the `x-krab:` block) by prefixing the
    child key. Values are de-quoted. This is intentionally a light parser — the
    frontmatter the compiler emits is simple and flat-ish, not arbitrary YAML.
    """
    pairs: List[Tuple[str, str]] = []
    parent = ""
    for raw in fm_lines:
        if not raw.strip():
            continue
        indent = len(raw) - len(raw.lstrip(" "))
        line = raw.strip()
        if ":" not in line:
            continue
        key, _, val = line.partition(":")
        key = key.strip()
        val = val.strip().strip('"').strip("'")
        if val == "" and indent == 0:
            # a parent key like "x-krab:" with children below it
            parent = key
            continue
        label = f"{parent}.{key}" if (indent > 0 and parent) else key
        if indent == 0:
            parent = ""  # reset once we're back at top level
        pairs.append((label, val))
    return pairs


# --------------------------------------------------------------------------- #
# Inline markdown (**bold**, `code`)
# --------------------------------------------------------------------------- #
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honoring **bold** and `code` spans."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


# --------------------------------------------------------------------------- #
# Styling setup
# --------------------------------------------------------------------------- #
def _setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # US Letter, 1" margins (the docx skill warns python's default is Letter on
    # python-docx, but we set it explicitly to be safe across viewers).
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(1))


# --------------------------------------------------------------------------- #
# Public entrypoint
# --------------------------------------------------------------------------- #
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def skill_md_to_docx(skill_md: str, *, slug: str = "skill") -> bytes:
    """Convert a compiled skill_md string into .docx bytes."""
    fm_lines, body = _split_frontmatter(skill_md)

    doc = Document()
    _setup_styles(doc)

    # --- metadata table from frontmatter -------------------------------- #
    pairs = _flatten_frontmatter(fm_lines)
    if pairs:
        heading = doc.add_paragraph()
        run = heading.add_run("Skill Metadata")
        run.bold = True
        run.font.size = Pt(13)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for key, val in pairs:
            row = table.add_row().cells
            kp = row[0].paragraphs[0]
            kr = kp.add_run(key)
            kr.bold = True
            row[1].paragraphs[0].add_run(val)
        doc.add_paragraph()  # spacer

    # --- body markdown -------------------------------------------------- #
    in_para_buffer: List[str] = []

    def flush_para():
        if in_para_buffer:
            text = " ".join(in_para_buffer).strip()
            if text:
                p = doc.add_paragraph()
                _add_inline_runs(p, text)
            in_para_buffer.clear()

    for raw_line in body.splitlines():
        line = raw_line.rstrip()

        if not line.strip():
            flush_para()
            continue

        m = HEADING_RE.match(line)
        if m:
            flush_para()
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            continue

        m = ORDERED_RE.match(line)
        if m:
            flush_para()
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m.group(2).strip())
            continue

        m = UNORDERED_RE.match(line)
        if m:
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1).strip())
            continue

        # plain text — accumulate into the current paragraph
        in_para_buffer.append(line.strip())

    flush_para()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
