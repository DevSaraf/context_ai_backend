"""
File Parser Module
Extracts text from uploaded files: PDF, DOCX, TXT, CSV, Markdown

Each parser returns a ParsedDocument with the extracted text and metadata.
The text then flows into chunking.py → embedding.py → database.

Dependencies to add to requirements.txt:
    pymupdf==1.25.3
    python-docx==1.1.2

Install:
    pip install pymupdf python-docx
"""

import os
import re
import csv
import io
import tempfile
from typing import Optional
from dataclasses import dataclass, field
from datetime import datetime


@dataclass
class ParsedDocument:
    """Result of parsing a file."""
    text: str                              # Extracted full text
    filename: str                          # Original filename
    file_type: str                         # pdf, docx, txt, csv, md
    page_count: int = 0                    # Number of pages (PDF only)
    word_count: int = 0                    # Total words extracted
    metadata: dict = field(default_factory=dict)  # Title, author, etc.
    error: Optional[str] = None            # Set if parsing failed

    @property
    def is_valid(self) -> bool:
        return self.error is None and len(self.text.strip()) > 0


# ============== MAIN ENTRY POINT ==============

async def parse_uploaded_file(file) -> ParsedDocument:
    """
    Parse an uploaded file (FastAPI UploadFile object).
    
    Usage in your endpoint:
        from app.file_parser import parse_uploaded_file

        @app.post("/knowledge/upload-file")
        async def upload_file(file: UploadFile, ...):
            parsed = await parse_uploaded_file(file)
            if not parsed.is_valid:
                raise HTTPException(400, parsed.error)
            chunks = chunk_text(parsed.text)
            ...
    """
    filename = file.filename or "unknown"
    extension = _get_extension(filename)
    content = await file.read()

    if len(content) == 0:
        return ParsedDocument(
            text="", filename=filename, file_type=extension,
            error="File is empty"
        )

    # Size limit: 20MB
    if len(content) > 20 * 1024 * 1024:
        return ParsedDocument(
            text="", filename=filename, file_type=extension,
            error="File too large. Maximum size is 20MB."
        )

    try:
        if extension == "pdf":
            return _parse_pdf(content, filename)
        elif extension == "docx":
            return _parse_docx(content, filename)
        elif extension == "txt":
            return _parse_txt(content, filename)
        elif extension == "md":
            return _parse_txt(content, filename)  # Markdown is just text
        elif extension == "csv":
            return _parse_csv(content, filename)
        else:
            return ParsedDocument(
                text="", filename=filename, file_type=extension,
                error=f"Unsupported file type: .{extension}. Supported: PDF, DOCX, TXT, CSV, MD"
            )
    except Exception as e:
        return ParsedDocument(
            text="", filename=filename, file_type=extension,
            error=f"Failed to parse file: {str(e)}"
        )


def parse_raw_text(text: str, source_name: str = "pasted_text") -> ParsedDocument:
    """Parse raw text content (for backward compatibility with your existing endpoint)."""
    cleaned = _clean_text(text)
    return ParsedDocument(
        text=cleaned,
        filename=source_name,
        file_type="text",
        word_count=len(cleaned.split()),
        metadata={"source": "direct_paste"}
    )


# ============== FILE TYPE PARSERS ==============

def _parse_pdf(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from PDF using PyMuPDF (fitz)."""
    import fitz  # pymupdf

    doc = fitz.open(stream=content, filetype="pdf")
    
    pages_text = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    full_text = _clean_text(full_text)

    # Extract metadata
    metadata = {}
    pdf_meta = doc.metadata
    if pdf_meta:
        if pdf_meta.get("title"):
            metadata["title"] = pdf_meta["title"]
        if pdf_meta.get("author"):
            metadata["author"] = pdf_meta["author"]
        if pdf_meta.get("subject"):
            metadata["subject"] = pdf_meta["subject"]

    doc.close()

    if not full_text.strip():
        return ParsedDocument(
            text="", filename=filename, file_type="pdf",
            page_count=len(doc),
            error="Could not extract text from PDF. It may be scanned/image-based."
        )

    return ParsedDocument(
        text=full_text,
        filename=filename,
        file_type="pdf",
        page_count=len(pages_text),
        word_count=len(full_text.split()),
        metadata=metadata
    )


def _parse_docx(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from DOCX using python-docx."""
    from docx import Document as DocxDocument

    # python-docx needs a file-like object
    doc = DocxDocument(io.BytesIO(content))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                    prefix = "#" * level_num
                    paragraphs.append(f"{prefix} {text}")
                except ValueError:
                    paragraphs.append(text)
            else:
                paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text:
            paragraphs.append(table_text)

    full_text = "\n\n".join(paragraphs)
    full_text = _clean_text(full_text)

    # Extract metadata
    metadata = {}
    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author

    return ParsedDocument(
        text=full_text,
        filename=filename,
        file_type="docx",
        word_count=len(full_text.split()),
        metadata=metadata
    )


def _parse_txt(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from plain text or markdown files."""
    # Try common encodings
    text = None
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            text = content.decode(encoding)
            break
        except (UnicodeDecodeError, ValueError):
            continue

    if text is None:
        return ParsedDocument(
            text="", filename=filename, file_type="txt",
            error="Could not decode file. Unsupported text encoding."
        )

    text = _clean_text(text)
    file_type = "md" if filename.lower().endswith(".md") else "txt"

    return ParsedDocument(
        text=text,
        filename=filename,
        file_type=file_type,
        word_count=len(text.split()),
        metadata={"encoding": "utf-8"}
    )


def _parse_csv(content: bytes, filename: str) -> ParsedDocument:
    """Convert CSV rows into readable text for embedding."""
    text_content = content.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text_content))

    rows_text = []
    row_count = 0
    for row in reader:
        # Convert each row into a natural language sentence
        parts = []
        for key, value in row.items():
            if key and value and value.strip():
                parts.append(f"{key}: {value.strip()}")
        if parts:
            rows_text.append(". ".join(parts) + ".")
            row_count += 1

        # Limit to 10,000 rows
        if row_count >= 10000:
            break

    full_text = "\n\n".join(rows_text)

    return ParsedDocument(
        text=full_text,
        filename=filename,
        file_type="csv",
        word_count=len(full_text.split()),
        metadata={"rows": row_count, "columns": list(reader.fieldnames or [])}
    )


# ============== HELPERS ==============

def _get_extension(filename: str) -> str:
    """Get lowercase file extension without the dot."""
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    return ext


def _clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove junk characters."""
    # Remove null bytes
    text = text.replace("\x00", "")
    
    # Normalize unicode whitespace
    text = re.sub(r'[\xa0\u200b\u200c\u200d\ufeff]', ' ', text)
    
    # Collapse multiple spaces (but preserve newlines)
    text = re.sub(r'[^\S\n]+', ' ', text)
    
    # Collapse 3+ newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove lines that are just dots, dashes, or underscores (decorative separators)
    text = re.sub(r'\n[.\-_=]{3,}\n', '\n\n', text)
    
    return text.strip()


def _extract_table_text(table) -> str:
    """Extract text from a DOCX table as readable rows."""
    rows_text = []
    headers = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        
        if i == 0:
            # First row is likely headers
            headers = cells
            continue

        if headers:
            # Format as "Header: Value" pairs
            parts = []
            for h, c in zip(headers, cells):
                if h and c:
                    parts.append(f"{h}: {c}")
            if parts:
                rows_text.append(". ".join(parts) + ".")
        else:
            # No headers, just join cells
            non_empty = [c for c in cells if c]
            if non_empty:
                rows_text.append(", ".join(non_empty))

    return "\n".join(rows_text)
